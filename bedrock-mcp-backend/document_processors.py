"""
📄 Document Processors - Generadores de DOCX y XLSX
Crea documentos profesionales para SOWs, calculadoras y reportes
"""

import os
import uuid
from datetime import datetime
from typing import Dict, List, Optional, Any
import logging

# Importaciones para DOCX
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Importaciones para XLSX
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    📄 Procesador base para documentos
    """
    
    def __init__(self):
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.unique_id = str(uuid.uuid4())[:8]
    
    def get_unique_filename(self, base_name: str, extension: str) -> str:
        """Genera nombre único para archivo"""
        return f"{base_name}_{self.timestamp}_{self.unique_id}.{extension}"

class SOWGenerator(DocumentProcessor):
    """
    📋 Generador de Statement of Work (SOW) en formato DOCX
    """
    
    def __init__(self):
        super().__init__()
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Ejecute: pip install python-docx")
    
    def generate_sow(self, project_data: Dict[str, Any]) -> bytes:
        """
        📋 Genera SOW completo en formato DOCX
        """
        doc = Document()
        
        # Configurar estilos
        self._setup_styles(doc)
        
        # 1. Encabezado
        self._add_header(doc, project_data)
        
        # 2. Objetivo y Alcances
        self._add_objectives_scope(doc, project_data)
        
        # 3. Información Técnica
        self._add_technical_info(doc, project_data)
        
        # 4. Plan de Actividades
        self._add_activity_plan(doc, project_data)
        
        # 5. Conclusiones
        self._add_conclusions(doc, project_data)
        
        # 6. Anexos
        self._add_appendices(doc, project_data)
        
        # Guardar en memoria
        from io import BytesIO
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
    
    def _setup_styles(self, doc):
        """Configurar estilos del documento"""
        styles = doc.styles
        
        # Estilo para títulos principales
        if 'Custom Heading 1' not in [s.name for s in styles]:
            heading_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = 'Arial'
            heading_style.font.size = Pt(16)
            heading_style.font.bold = True
            heading_style.paragraph_format.space_after = Pt(12)
    
    def _add_header(self, doc, project_data):
        """Agregar encabezado del documento"""
        # Título principal
        title = doc.add_heading('STATEMENT OF WORK (SOW)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del proyecto
        doc.add_paragraph()
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Proyecto:', project_data.get('project_name', 'Proyecto AWS')),
            ('Cliente:', project_data.get('client_name', 'Cliente')),
            ('Fecha:', datetime.now().strftime('%d/%m/%Y')),
            ('Versión:', project_data.get('version', '1.0')),
            ('Responsable:', project_data.get('responsible', 'AWS Solutions Architect'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
            info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    def _add_objectives_scope(self, doc, project_data):
        """Agregar objetivos y alcances"""
        doc.add_heading('1. OBJETIVO Y ALCANCES', level=1)
        
        # Objetivo
        doc.add_heading('1.1 Objetivo General', level=2)
        objective = project_data.get('objective', 
            'Implementar una solución en AWS que permita optimizar la infraestructura, '
            'mejorar la escalabilidad y reducir costos operativos.')
        doc.add_paragraph(objective)
        
        # Alcances
        doc.add_heading('1.2 Alcances del Proyecto', level=2)
        scopes = project_data.get('scopes', [
            'Análisis de la infraestructura actual',
            'Diseño de arquitectura AWS optimizada',
            'Implementación de servicios AWS core',
            'Configuración de monitoreo y alertas',
            'Documentación técnica completa',
            'Transferencia de conocimiento al equipo'
        ])
        
        for scope in scopes:
            doc.add_paragraph(scope, style='List Bullet')
        
        # Fuera de alcance
        doc.add_heading('1.3 Fuera de Alcance', level=2)
        out_of_scope = project_data.get('out_of_scope', [
            'Migración de datos legacy no especificados',
            'Desarrollo de aplicaciones custom',
            'Soporte post-implementación (se cotiza por separado)'
        ])
        
        for item in out_of_scope:
            doc.add_paragraph(item, style='List Bullet')
    
    def _add_technical_info(self, doc, project_data):
        """Agregar información técnica"""
        doc.add_heading('2. INFORMACIÓN TÉCNICA', level=1)
        
        # Arquitectura propuesta
        doc.add_heading('2.1 Arquitectura Propuesta', level=2)
        architecture = project_data.get('architecture', 
            'La solución propuesta utiliza servicios AWS nativos para garantizar '
            'alta disponibilidad, escalabilidad y seguridad.')
        doc.add_paragraph(architecture)
        
        # Servicios AWS
        doc.add_heading('2.2 Servicios AWS Utilizados', level=2)
        services = project_data.get('aws_services', [
            'Amazon EC2 - Instancias de cómputo',
            'Amazon RDS - Base de datos relacional',
            'Amazon S3 - Almacenamiento de objetos',
            'Amazon CloudFront - CDN global',
            'AWS Lambda - Funciones serverless',
            'Amazon CloudWatch - Monitoreo y alertas'
        ])
        
        for service in services:
            doc.add_paragraph(service, style='List Bullet')
        
        # Estimación de costos
        doc.add_heading('2.3 Estimación de Costos', level=2)
        costs = project_data.get('cost_estimation', {})
        
        if costs:
            cost_table = doc.add_table(rows=len(costs) + 1, cols=3)
            cost_table.style = 'Table Grid'
            
            # Headers
            headers = ['Servicio', 'Costo Mensual (USD)', 'Descripción']
            for i, header in enumerate(headers):
                cell = cost_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data
            for i, (service, details) in enumerate(costs.items(), 1):
                cost_table.cell(i, 0).text = service
                cost_table.cell(i, 1).text = f"${details.get('monthly_cost', 0):,.2f}"
                cost_table.cell(i, 2).text = details.get('description', '')
    
    def _add_activity_plan(self, doc, project_data):
        """Agregar plan de actividades"""
        doc.add_heading('3. PLAN DE ACTIVIDADES DE IMPLEMENTACIÓN', level=1)
        
        activities = project_data.get('activities', [
            {
                'phase': 'Fase 1: Análisis y Diseño',
                'duration': '2 semanas',
                'tasks': [
                    'Análisis de requerimientos',
                    'Diseño de arquitectura',
                    'Documentación técnica inicial'
                ]
            },
            {
                'phase': 'Fase 2: Implementación Core',
                'duration': '3 semanas',
                'tasks': [
                    'Configuración de VPC y redes',
                    'Despliegue de servicios principales',
                    'Configuración de seguridad'
                ]
            },
            {
                'phase': 'Fase 3: Testing y Optimización',
                'duration': '1 semana',
                'tasks': [
                    'Pruebas de funcionalidad',
                    'Optimización de performance',
                    'Configuración de monitoreo'
                ]
            },
            {
                'phase': 'Fase 4: Go-Live y Documentación',
                'duration': '1 semana',
                'tasks': [
                    'Migración a producción',
                    'Documentación final',
                    'Transferencia de conocimiento'
                ]
            }
        ])
        
        for activity in activities:
            doc.add_heading(f"{activity['phase']} ({activity['duration']})", level=2)
            for task in activity['tasks']:
                doc.add_paragraph(task, style='List Bullet')
    
    def _add_conclusions(self, doc, project_data):
        """Agregar conclusiones"""
        doc.add_heading('4. CONCLUSIONES', level=1)
        
        conclusions = project_data.get('conclusions', [
            'La implementación propuesta permitirá una reducción estimada del 30-40% en costos operativos.',
            'La arquitectura AWS nativa garantiza alta disponibilidad y escalabilidad automática.',
            'El proyecto se completará en aproximadamente 7 semanas con entregables claros en cada fase.',
            'Se incluye documentación completa y transferencia de conocimiento para el equipo interno.'
        ])
        
        for conclusion in conclusions:
            doc.add_paragraph(conclusion, style='List Bullet')
    
    def _add_appendices(self, doc, project_data):
        """Agregar anexos"""
        doc.add_heading('5. ANEXOS', level=1)
        
        appendices = project_data.get('appendices', [
            'Diagramas de arquitectura detallados',
            'Matriz de responsabilidades RACI',
            'Plan de contingencia y rollback',
            'Checklist de seguridad AWS'
        ])
        
        for appendix in appendices:
            doc.add_paragraph(appendix, style='List Bullet')

class AWSCalculatorGenerator(DocumentProcessor):
    """
    🧮 Generador de Calculadora AWS en formato XLSX
    """
    
    def __init__(self):
        super().__init__()
        if not XLSX_AVAILABLE:
            raise ImportError("openpyxl no está instalado. Ejecute: pip install openpyxl")
    
    def generate_calculator(self, calculator_data: Dict[str, Any]) -> bytes:
        """
        🧮 Genera calculadora AWS completa en formato XLSX
        """
        wb = Workbook()
        
        # Eliminar hoja por defecto
        wb.remove(wb.active)
        
        # Crear hojas
        self._create_summary_sheet(wb, calculator_data)
        self._create_compute_sheet(wb, calculator_data)
        self._create_storage_sheet(wb, calculator_data)
        self._create_database_sheet(wb, calculator_data)
        self._create_networking_sheet(wb, calculator_data)
        
        # Guardar en memoria
        from io import BytesIO
        wb_bytes = BytesIO()
        wb.save(wb_bytes)
        wb_bytes.seek(0)
        
        return wb_bytes.getvalue()
    
    def _create_summary_sheet(self, wb, data):
        """Crear hoja de resumen"""
        ws = wb.create_sheet("Resumen", 0)
        
        # Título
        ws['A1'] = 'CALCULADORA DE COSTOS AWS'
        ws['A1'].font = Font(size=16, bold=True)
        ws.merge_cells('A1:F1')
        
        # Información del proyecto
        ws['A3'] = 'Proyecto:'
        ws['B3'] = data.get('project_name', 'Proyecto AWS')
        ws['A4'] = 'Fecha:'
        ws['B4'] = datetime.now().strftime('%d/%m/%Y')
        ws['A5'] = 'Región:'
        ws['B5'] = data.get('region', 'us-east-1')
        
        # Headers de resumen
        headers = ['Categoría', 'Servicio', 'Cantidad', 'Costo Unitario', 'Costo Mensual', 'Costo Anual']
        for i, header in enumerate(headers, 1):
            cell = ws.cell(row=7, column=i)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.font = Font(color="FFFFFF", bold=True)
        
        # Datos de ejemplo
        sample_data = [
            ['Compute', 'EC2 t3.medium', '2', '$0.0416/hora', '$59.90', '$718.80'],
            ['Storage', 'S3 Standard', '100 GB', '$0.023/GB', '$2.30', '$27.60'],
            ['Database', 'RDS MySQL t3.micro', '1', '$0.017/hora', '$12.24', '$146.88'],
            ['Networking', 'CloudFront', '1 TB', '$0.085/GB', '$85.00', '$1,020.00']
        ]
        
        for i, row_data in enumerate(sample_data, 8):
            for j, value in enumerate(row_data, 1):
                ws.cell(row=i, column=j).value = value
        
        # Total
        total_row = len(sample_data) + 9
        ws.cell(row=total_row, column=5).value = 'TOTAL MENSUAL:'
        ws.cell(row=total_row, column=5).font = Font(bold=True)
        ws.cell(row=total_row, column=6).value = '$159.44'
        ws.cell(row=total_row, column=6).font = Font(bold=True)
        
        # Ajustar columnas
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
    
    def _create_compute_sheet(self, wb, data):
        """Crear hoja de cómputo"""
        ws = wb.create_sheet("Compute")
        
        ws['A1'] = 'SERVICIOS DE CÓMPUTO'
        ws['A1'].font = Font(size=14, bold=True)
        
        # EC2 Calculator
        ws['A3'] = 'Amazon EC2'
        ws['A3'].font = Font(bold=True)
        
        headers = ['Tipo de Instancia', 'vCPUs', 'RAM (GB)', 'Precio/Hora', 'Horas/Mes', 'Costo Mensual']
        for i, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=i)
            cell.value = header
            cell.font = Font(bold=True)
        
        # Datos de instancias EC2
        ec2_data = [
            ['t3.nano', '2', '0.5', '$0.0052', '730', '=D5*E5'],
            ['t3.micro', '2', '1', '$0.0104', '730', '=D6*E6'],
            ['t3.small', '2', '2', '$0.0208', '730', '=D7*E7'],
            ['t3.medium', '2', '4', '$0.0416', '730', '=D8*E8'],
            ['t3.large', '2', '8', '$0.0832', '730', '=D9*E9']
        ]
        
        for i, row_data in enumerate(ec2_data, 5):
            for j, value in enumerate(row_data, 1):
                ws.cell(row=i, column=j).value = value
    
    def _create_storage_sheet(self, wb, data):
        """Crear hoja de almacenamiento"""
        ws = wb.create_sheet("Storage")
        
        ws['A1'] = 'SERVICIOS DE ALMACENAMIENTO'
        ws['A1'].font = Font(size=14, bold=True)
        
        # S3 Calculator
        ws['A3'] = 'Amazon S3'
        ws['A3'].font = Font(bold=True)
        
        headers = ['Clase de Almacenamiento', 'Precio/GB/Mes', 'GB Almacenados', 'Costo Mensual']
        for i, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=i)
            cell.value = header
            cell.font = Font(bold=True)
        
        s3_data = [
            ['Standard', '$0.023', '0', '=B5*C5'],
            ['Standard-IA', '$0.0125', '0', '=B6*C6'],
            ['Glacier', '$0.004', '0', '=B7*C7'],
            ['Glacier Deep Archive', '$0.00099', '0', '=B8*C8']
        ]
        
        for i, row_data in enumerate(s3_data, 5):
            for j, value in enumerate(row_data, 1):
                ws.cell(row=i, column=j).value = value
    
    def _create_database_sheet(self, wb, data):
        """Crear hoja de bases de datos"""
        ws = wb.create_sheet("Database")
        
        ws['A1'] = 'SERVICIOS DE BASE DE DATOS'
        ws['A1'].font = Font(size=14, bold=True)
        
        # RDS Calculator
        ws['A3'] = 'Amazon RDS'
        ws['A3'].font = Font(bold=True)
        
        headers = ['Motor', 'Tipo de Instancia', 'Precio/Hora', 'Horas/Mes', 'Costo Mensual']
        for i, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=i)
            cell.value = header
            cell.font = Font(bold=True)
        
        rds_data = [
            ['MySQL', 'db.t3.micro', '$0.017', '730', '=C5*D5'],
            ['PostgreSQL', 'db.t3.small', '$0.034', '730', '=C6*D6'],
            ['Oracle', 'db.t3.medium', '$0.068', '730', '=C7*D7']
        ]
        
        for i, row_data in enumerate(rds_data, 5):
            for j, value in enumerate(row_data, 1):
                ws.cell(row=i, column=j).value = value
    
    def _create_networking_sheet(self, wb, data):
        """Crear hoja de networking"""
        ws = wb.create_sheet("Networking")
        
        ws['A1'] = 'SERVICIOS DE RED'
        ws['A1'].font = Font(size=14, bold=True)
        
        # CloudFront Calculator
        ws['A3'] = 'Amazon CloudFront'
        ws['A3'].font = Font(bold=True)
        
        headers = ['Región', 'Precio/GB', 'GB Transferidos', 'Costo Mensual']
        for i, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=i)
            cell.value = header
            cell.font = Font(bold=True)
        
        cf_data = [
            ['Estados Unidos', '$0.085', '0', '=B5*C5'],
            ['Europa', '$0.085', '0', '=B6*C6'],
            ['Asia Pacífico', '$0.140', '0', '=B7*C7']
        ]
        
        for i, row_data in enumerate(cf_data, 5):
            for j, value in enumerate(row_data, 1):
                ws.cell(row=i, column=j).value = value

# Funciones de utilidad
def generate_sow_document(project_data: Dict[str, Any]) -> bytes:
    """
    📋 Genera documento SOW
    """
    try:
        generator = SOWGenerator()
        return generator.generate_sow(project_data)
    except ImportError as e:
        logger.error(f"❌ Error generando SOW: {str(e)}")
        raise
    except Exception as e:
        logger.error(f"💥 Error inesperado generando SOW: {str(e)}")
        raise

def generate_aws_calculator(calculator_data: Dict[str, Any]) -> bytes:
    """
    🧮 Genera calculadora AWS
    """
    try:
        generator = AWSCalculatorGenerator()
        return generator.generate_calculator(calculator_data)
    except ImportError as e:
        logger.error(f"❌ Error generando calculadora: {str(e)}")
        raise
    except Exception as e:
        logger.error(f"💥 Error inesperado generando calculadora: {str(e)}")
        raise
